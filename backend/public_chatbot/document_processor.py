"""
Document Processor for Public Chatbot - File Upload & Format Conversion
Supports PDF, Word, HTML, TXT, CSV, JSON and other formats
Integrates with existing PublicKnowledgeDocument workflow
"""
import os
import logging
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
from datetime import datetime
import tempfile

# Core document processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# Django imports
from django.core.files.uploadedfile import UploadedFile
from django.core.exceptions import ValidationError

logger = logging.getLogger('public_chatbot')


class DocumentProcessor:
    """
    Process uploaded documents and convert them to text for PublicKnowledgeDocument
    """
    
    # Supported file types and their processors
    SUPPORTED_FORMATS = {
        # Text formats
        '.txt': 'process_text',
        '.md': 'process_markdown',
        '.markdown': 'process_markdown',
        
        # Microsoft Office
        '.docx': 'process_docx',
        '.doc': 'process_docx',  # Will attempt docx processing
        
        # PDF
        '.pdf': 'process_pdf',
        
        # Web formats
        '.html': 'process_html',
        '.htm': 'process_html',
        
        # Data formats
        '.csv': 'process_csv',
        '.json': 'process_json',
        
        # Excel (optional)
        '.xlsx': 'process_excel',
        '.xls': 'process_excel',
    }
    
    # File size limits (in bytes)
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_TOTAL_SIZE = 200 * 1024 * 1024  # 200MB for batch uploads
    
    def __init__(self):
        self.processed_docs = []
        self.errors = []
        self.warnings = []
    
    def process_uploaded_files(self, files: List[UploadedFile], 
                             default_category: str = 'general',
                             created_by: str = 'admin',
                             enable_security_scan: bool = True) -> Dict[str, Any]:
        """
        Process multiple uploaded files and return document data
        
        Args:
            files: List of Django UploadedFile objects
            default_category: Default category for documents
            created_by: Admin username
            enable_security_scan: Enable security validation
            
        Returns:
            Dict with processed documents, errors, and summary
        """
        logger.info(f"🔄 UPLOAD: Processing {len(files)} uploaded files")
        
        # Reset state
        self.processed_docs = []
        self.errors = []
        self.warnings = []
        
        # Security validation
        if enable_security_scan:
            from .security import DocumentSecurityValidator
            security_validator = DocumentSecurityValidator()
            security_result = security_validator.validate_upload_batch(files)
            
            if not security_result['valid']:
                # Add security errors
                for error in security_result['batch_errors']:
                    self.errors.append({
                        'file': 'batch',
                        'error': f"Security: {error['message']}"
                    })
                
                # Add file-specific security errors
                for filename, file_result in security_result['file_results'].items():
                    for error in file_result['errors']:
                        self.errors.append({
                            'file': filename,
                            'error': f"Security: {error}"
                        })
                
                # Add security warnings
                for warning in security_result['security_warnings']:
                    self.warnings.append({
                        'file': warning['file'],
                        'warning': f"Security: {warning['message']}"
                    })
                
                # If there are critical security errors, stop processing
                if security_result['files_invalid'] > 0:
                    logger.warning(f"🔒 SECURITY: {security_result['files_invalid']} files failed security validation")
                    return self._build_result()
        
        # Validate batch size and total size
        total_size = sum(f.size for f in files)
        if total_size > self.MAX_TOTAL_SIZE:
            self.errors.append({
                'file': 'batch',
                'error': f'Total upload size ({total_size/1024/1024:.1f}MB) exceeds limit ({self.MAX_TOTAL_SIZE/1024/1024}MB)'
            })
            return self._build_result()
        
        # Process each file
        for file_obj in files:
            try:
                self._process_single_file(file_obj, default_category, created_by)
            except Exception as e:
                logger.error(f"❌ UPLOAD: Error processing {file_obj.name}: {e}")
                self.errors.append({
                    'file': file_obj.name,
                    'error': str(e)
                })
        
        logger.info(f"✅ UPLOAD: Processed {len(self.processed_docs)} documents successfully, {len(self.errors)} errors")
        return self._build_result()
    
    def _process_single_file(self, file_obj: UploadedFile, category: str, created_by: str):
        """Process a single uploaded file"""
        # Validate file
        self._validate_file(file_obj)
        
        # Get file extension
        file_path = Path(file_obj.name)
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_FORMATS:
            raise ValidationError(f"Unsupported file format: {extension}")
        
        # Extract text content
        processor_method = self.SUPPORTED_FORMATS[extension]
        content = getattr(self, processor_method)(file_obj)
        
        if not content or len(content.strip()) < 10:
            self.warnings.append({
                'file': file_obj.name,
                'warning': 'Document appears to be empty or very short'
            })
            return
        
        # Create document data
        doc_data = {
            'title': self._extract_title(file_path.stem, content),
            'content': content,
            'category': category,
            'subcategory': self._detect_subcategory(content, extension),
            'source_url': f'upload://{file_obj.name}',
            'tags': self._extract_tags(content, file_obj.name),
            'created_by': created_by,
            'language': 'en',  # Could be enhanced with language detection
            'quality_score': self._calculate_quality_score(content),
            # File metadata
            'file_name': file_obj.name,
            'file_size': file_obj.size,
            'file_type': extension,
            'content_hash': hashlib.sha256(content.encode()).hexdigest()
        }
        
        self.processed_docs.append(doc_data)
        logger.info(f"📄 UPLOAD: Processed {file_obj.name} -> {len(content)} chars")
    
    def _validate_file(self, file_obj: UploadedFile):
        """Validate uploaded file for security and size"""
        # Size check
        if file_obj.size > self.MAX_FILE_SIZE:
            raise ValidationError(f"File too large: {file_obj.size/1024/1024:.1f}MB (max: {self.MAX_FILE_SIZE/1024/1024}MB)")
        
        # Name validation (security)
        if not file_obj.name or '..' in file_obj.name or '/' in file_obj.name:
            raise ValidationError("Invalid file name")
        
        # Content type validation
        content_type = file_obj.content_type
        if content_type and 'script' in content_type.lower():
            raise ValidationError("Script files not allowed")
        
        logger.debug(f"🔍 UPLOAD: Validated {file_obj.name} ({file_obj.size} bytes, {content_type})")
    
    # Format-specific processors
    
    def process_text(self, file_obj: UploadedFile) -> str:
        """Process plain text files"""
        try:
            # Try UTF-8 first, fallback to other encodings
            for encoding in ['utf-8', 'latin1', 'cp1252']:
                try:
                    content = file_obj.read().decode(encoding)
                    file_obj.seek(0)  # Reset for potential retry
                    return content.strip()
                except UnicodeDecodeError:
                    file_obj.seek(0)
                    continue
            
            raise ValidationError("Could not decode text file with any supported encoding")
        except Exception as e:
            raise ValidationError(f"Error reading text file: {e}")
    
    def process_markdown(self, file_obj: UploadedFile) -> str:
        """Process Markdown files"""
        try:
            raw_content = self.process_text(file_obj)
            
            if MARKDOWN_AVAILABLE:
                # Convert markdown to plain text (remove markdown syntax)
                html = markdown.markdown(raw_content)
                if HTML_AVAILABLE:
                    soup = BeautifulSoup(html, 'html.parser')
                    return soup.get_text('\n').strip()
                else:
                    # Basic markdown removal
                    import re
                    # Remove common markdown syntax
                    text = re.sub(r'#{1,6}\s+', '', raw_content)  # Headers
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
                    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
                    text = re.sub(r'`(.*?)`', r'\1', text)  # Code
                    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
                    return text.strip()
            else:
                return raw_content
                
        except Exception as e:
            raise ValidationError(f"Error processing Markdown file: {e}")
    
    def process_pdf(self, file_obj: UploadedFile) -> str:
        """Process PDF files"""
        if not PDF_AVAILABLE:
            raise ValidationError("PDF processing not available. Install PyPDF2 and pdfplumber.")
        
        try:
            content_parts = []
            
            # Try pdfplumber first (better text extraction)
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(file_obj.read())
                    tmp_file.flush()
                    
                    with pdfplumber.open(tmp_file.name) as pdf:
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                content_parts.append(text)
                
                os.unlink(tmp_file.name)
                
            except Exception:
                # Fallback to PyPDF2
                file_obj.seek(0)
                reader = PyPDF2.PdfReader(file_obj)
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
            
            content = '\n\n'.join(content_parts).strip()
            if not content:
                raise ValidationError("Could not extract text from PDF")
            
            return content
            
        except Exception as e:
            raise ValidationError(f"Error processing PDF: {e}")
    
    def process_docx(self, file_obj: UploadedFile) -> str:
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            raise ValidationError("Word processing not available. Install python-docx.")
        
        try:
            # Save to temporary file (python-docx needs file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_obj.read())
                tmp_file.flush()
                
                doc = DocxDocument(tmp_file.name)
                content_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text.strip())
                
                # Extract table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            content_parts.append(row_text)
            
            os.unlink(tmp_file.name)
            
            content = '\n\n'.join(content_parts).strip()
            if not content:
                raise ValidationError("Could not extract text from Word document")
            
            return content
            
        except Exception as e:
            raise ValidationError(f"Error processing Word document: {e}")
    
    def process_html(self, file_obj: UploadedFile) -> str:
        """Process HTML files"""
        if not HTML_AVAILABLE:
            raise ValidationError("HTML processing not available. Install beautifulsoup4.")
        
        try:
            html_content = self.process_text(file_obj)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text('\n')
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            content = '\n'.join(line for line in lines if line)
            
            return content
            
        except Exception as e:
            raise ValidationError(f"Error processing HTML file: {e}")
    
    def process_csv(self, file_obj: UploadedFile) -> str:
        """Process CSV files"""
        import csv
        try:
            content = self.process_text(file_obj)
            
            # Parse CSV and convert to readable format
            lines = content.splitlines()
            reader = csv.reader(lines)
            
            content_parts = []
            headers = next(reader, None)
            
            if headers:
                content_parts.append("Data columns: " + " | ".join(headers))
                content_parts.append("-" * 50)
                
                for i, row in enumerate(reader):
                    if i >= 100:  # Limit to first 100 rows
                        content_parts.append(f"... and {sum(1 for _ in reader)} more rows")
                        break
                    
                    if len(row) == len(headers):
                        row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row))
                        content_parts.append(row_text)
            
            return '\n'.join(content_parts)
            
        except Exception as e:
            raise ValidationError(f"Error processing CSV file: {e}")
    
    def process_json(self, file_obj: UploadedFile) -> str:
        """Process JSON files"""
        import json
        try:
            content = self.process_text(file_obj)
            data = json.loads(content)
            
            # Convert JSON to readable text format
            def json_to_text(obj, level=0):
                indent = "  " * level
                if isinstance(obj, dict):
                    parts = [f"{indent}{k}: {json_to_text(v, level+1)}" for k, v in obj.items()]
                    return "\n".join(parts)
                elif isinstance(obj, list):
                    if len(obj) > 10:  # Limit large arrays
                        parts = [f"{indent}- {json_to_text(item, level+1)}" for item in obj[:10]]
                        parts.append(f"{indent}... and {len(obj)-10} more items")
                        return "\n".join(parts)
                    else:
                        parts = [f"{indent}- {json_to_text(item, level+1)}" for item in obj]
                        return "\n".join(parts)
                else:
                    return str(obj)
            
            return json_to_text(data)
            
        except Exception as e:
            raise ValidationError(f"Error processing JSON file: {e}")
    
    def process_excel(self, file_obj: UploadedFile) -> str:
        """Process Excel files"""
        if not EXCEL_AVAILABLE:
            raise ValidationError("Excel processing not available. Install openpyxl.")
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(file_obj.read())
                tmp_file.flush()
                
                workbook = openpyxl.load_workbook(tmp_file.name, data_only=True)
                content_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    content_parts.append(f"Sheet: {sheet_name}")
                    content_parts.append("-" * 30)
                    
                    # Get data from sheet
                    rows = list(sheet.iter_rows(values_only=True))
                    if rows:
                        # Headers
                        headers = [str(cell) if cell is not None else "" for cell in rows[0]]
                        content_parts.append(" | ".join(headers))
                        
                        # Data rows (limit to 50)
                        for row in rows[1:51]:
                            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                            if row_text.strip():
                                content_parts.append(row_text)
                        
                        if len(rows) > 51:
                            content_parts.append(f"... and {len(rows)-51} more rows")
                    
                    content_parts.append("")  # Blank line between sheets
            
            os.unlink(tmp_file.name)
            return '\n'.join(content_parts)
            
        except Exception as e:
            raise ValidationError(f"Error processing Excel file: {e}")
    
    # Helper methods
    
    def _extract_title(self, filename: str, content: str) -> str:
        """Extract meaningful title from filename and content"""
        # Clean filename
        title = filename.replace('_', ' ').replace('-', ' ')
        title = ' '.join(word.capitalize() for word in title.split())
        
        # Try to find a better title in content (first line or heading)
        lines = content.split('\n')[:5]  # Check first 5 lines
        for line in lines:
            line = line.strip()
            if line and len(line) < 100 and len(line) > 5:
                # Check if it looks like a title
                if not line.endswith('.') or line.count(' ') <= 10:
                    return line
        
        return title if title else "Untitled Document"
    
    def _detect_subcategory(self, content: str, extension: str) -> str:
        """Detect subcategory based on content analysis"""
        content_lower = content.lower()
        
        # Technical content
        if any(term in content_lower for term in ['api', 'code', 'function', 'class', 'method', 'algorithm']):
            return 'technical'
        
        # Documentation
        if any(term in content_lower for term in ['guide', 'manual', 'documentation', 'instructions', 'tutorial']):
            return 'documentation'
        
        # FAQ
        if 'question' in content_lower and 'answer' in content_lower:
            return 'faq'
        
        # Data/Analytics
        if extension in ['.csv', '.xlsx', '.json'] or 'data' in content_lower:
            return 'data'
        
        return 'general'
    
    def _extract_tags(self, content: str, filename: str) -> str:
        """Extract relevant tags from content and filename"""
        tags = set()
        
        # Add file extension as tag
        ext = Path(filename).suffix.lower()
        if ext:
            tags.add(ext[1:])  # Remove dot
        
        # Extract keywords from filename
        name_parts = Path(filename).stem.lower().replace('_', ' ').replace('-', ' ').split()
        tags.update(part for part in name_parts if len(part) > 3)
        
        # Simple keyword extraction from content
        content_words = content.lower().split()
        common_keywords = ['api', 'guide', 'tutorial', 'documentation', 'manual', 'faq', 'data', 'analysis']
        
        for keyword in common_keywords:
            if keyword in content_words:
                tags.add(keyword)
        
        return ', '.join(sorted(tags)[:10])  # Limit to 10 tags
    
    def _calculate_quality_score(self, content: str) -> int:
        """Calculate quality score based on content characteristics"""
        score = 50  # Base score
        
        # Length factor
        length = len(content)
        if length > 5000:
            score += 20
        elif length > 1000:
            score += 10
        elif length < 100:
            score -= 20
        
        # Structure indicators
        if '\n\n' in content:  # Has paragraphs
            score += 10
        
        if any(marker in content for marker in ['1.', '2.', '-', '*']):  # Has lists
            score += 10
        
        # Content quality indicators
        sentences = content.count('.') + content.count('!') + content.count('?')
        if sentences > 10:
            score += 10
        
        # Cap score
        return min(max(score, 0), 100)
    
    def _build_result(self) -> Dict[str, Any]:
        """Build final result dictionary"""
        return {
            'success': len(self.errors) == 0,
            'processed_count': len(self.processed_docs),
            'error_count': len(self.errors),
            'warning_count': len(self.warnings),
            'documents': self.processed_docs,
            'errors': self.errors,
            'warnings': self.warnings,
            'summary': {
                'total_files': len(self.processed_docs) + len(self.errors),
                'successful': len(self.processed_docs),
                'failed': len(self.errors),
                'formats_processed': list(set(doc.get('file_type', 'unknown') for doc in self.processed_docs))
            }
        }
    
    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported file formats"""
        return list(cls.SUPPORTED_FORMATS.keys())
    
    @classmethod
    def check_dependencies(cls) -> Dict[str, bool]:
        """Check which format processors are available"""
        return {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'excel': EXCEL_AVAILABLE,
            'html': HTML_AVAILABLE,
            'markdown': MARKDOWN_AVAILABLE
        }